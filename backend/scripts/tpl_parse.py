#!/usr/bin/env python3
"""tpl_parse.py — 解析 docx 模板结构,输出段落与表格 JSON。

用法(由 Go 后端以子进程调用):
    tpl_parse.py <docx_path>
stdin 不接收数据;结构 JSON 输出到 stdout,错误信息输出到 stderr 并退出码非 0。
输出格式:
    {"paragraphs": [{"idx": 0, "text": "...", "style": "Heading 1"}, ...],
     "tables": [["c1", "c2", "c3"], ...]}
仅输出非空文本的段落,保留原始顺序(段落 idx 为文档内真实索引)。
"""
import json
import sys

import docx

# Windows 下 stdout 默认 GBK 编码,Go 子进程以 UTF-8 交换数据,必须统一
sys.stdout.reconfigure(encoding="utf-8")


def main() -> int:
    if len(sys.argv) != 2:
        print("usage: tpl_parse.py <docx_path>", file=sys.stderr)
        return 2
    path = sys.argv[1]
    try:
        document = docx.Document(path)
    except Exception as exc:  # 文件损坏或非 docx
        print(f"cannot open docx: {exc}", file=sys.stderr)
        return 1

    paragraphs = []
    for idx, para in enumerate(document.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style is not None else ""
        paragraphs.append({"idx": idx, "text": text, "style": style})

    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)

    print(json.dumps({"paragraphs": paragraphs, "tables": tables}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    sys.exit(main())
