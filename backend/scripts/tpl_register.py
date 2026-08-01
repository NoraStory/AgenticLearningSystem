#!/usr/bin/env python3
"""tpl_register.py — 把用户上传的无占位符 docx 模板注入 Jinja2 占位符,生成已注册模板。

用法(由 Go 后端以子进程调用):
    tpl_register.py <docx_path> <out_path>
章节结构 JSON 从 stdin 读取:
    {"sections": [{"title": "基本信息", "items": ["姓名", "电话"]}, ...]}
其中 items 为 LLM/用户确认后的条目文本列表;每个 items 条目文本(去掉首尾空白后)
若能在对应 section 的段落中被找到,则把该条目文本替换为 {{ sections[i].items[j] }}。
未匹配的条目仍以普通文本保留(供用户核对,不报错)。

输出:注入后的 docx 写到 out_path;stdout 输出 {"written": "<out_path>"}。
"""
import json
import sys

import docx

# Windows 下 stdin/stdout 默认 GBK 编码,Go 子进程以 UTF-8 交换数据,必须统一
sys.stdin.reconfigure(encoding="utf-8")
sys.stdout.reconfigure(encoding="utf-8")


def section_paragraphs(paras, start, end):
    """返回段落列表下标 [start, end) 中非空的 (下标, 文本)。"""
    return [(i, p.text.strip()) for i, p in enumerate(paras) if start <= i < end and p.text.strip()]


def run_marker(para, text):
    """把 para 的文本设置为 text(合并为单 run,解决 Word 跨 run 拆占位符问题)。"""
    para.text = ""  # 清空全部 run
    run = para.add_run(text)
    return run


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: tpl_register.py <docx_path> <out_path>", file=sys.stderr)
        return 2
    src, out = sys.argv[1], sys.argv[2]
    try:
        payload = json.load(sys.stdin)
    except Exception as exc:
        print(f"invalid stdin json: {exc}", file=sys.stderr)
        return 2
    sections = payload.get("sections") or []
    if not sections:
        print("sections must not be empty", file=sys.stderr)
        return 2

    try:
        document = docx.Document(src)
    except Exception as exc:
        print(f"cannot open docx: {exc}", file=sys.stderr)
        return 1

    paras = document.paragraphs
    n = len(paras)
    for si, section in enumerate(sections):
        title = (section.get("title") or "").strip()
        items = [it.strip() for it in (section.get("items") or []) if it.strip()]
        if not items:
            continue
        # 按段落顺序扫描,找 title 段之后的 items 条目
        start = 0
        for i, para in enumerate(paras):
            if para.text.strip() == title:
                start = i + 1
                break
        remaining = list(items)
        for i, para in enumerate(paras):
            if i < start:
                continue
            ptext = para.text.strip()
            if not ptext:
                continue
            for item in remaining:
                if ptext == item:
                    # Jinja2 嵌套访问必须用方括号语法;点号索引(sections.0.items.0)会静默渲染为空
                    marker = "{{ sections[%d]['items'][%d] }}" % (si, items.index(item))
                    run_marker(para, marker)
                    remaining.remove(item)
                    break
            if not remaining:
                break

    document.save(out)
    print(json.dumps({"written": out}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    sys.exit(main())
